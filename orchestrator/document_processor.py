"""Universal document processor supporting PDF, DOCX, DOC, TXT, and more."""

import re
from typing import List, Dict, Any
from pathlib import Path


class DocumentProcessor:
    """Process various document formats and extract clean content."""

    # Section headers that mark beginning of unwanted large sections
    # These will remove everything from this point forward to next major section
    UNWANTED_SECTION_HEADERS = [
        r"^table of contents$",
        r"^list of figures$",
        r"^list of tables$",
        r"^preface$",
        r"^foreword$",
        r"^references$",
        r"^bibliography$",
        r"^works cited$",
        r"^appendix",
        r"^appendices$",
        r"^endnotes$",
        r"^footnotes$",
        r"^glossary$",
        r"^index$",
        r"^acknowledgments?$",
        r"^about the author$",
    ]

    # Unwanted phrases that appear in content (remove entire paragraph if found)
    UNWANTED_PHRASES = [
        # Boilerplate/meta content
        r"^copyright",
        r"©",
        r"all rights reserved",
        r"^isbn",
        r"published by",
        r"publication date",

        # Introductory fluff
        r"this (book|document|guide|manual|chapter) (will|shall|aims to|is designed to)",
        r"in this (chapter|section), (we|you) will",
        r"about this (book|document|guide|manual)",

        # Contact/promotional
        r"for more information, (visit|see|contact)",
        r"visit our website",
        r"www\.",
        r"http://",
        r"https://",
        r"email:",
        r"contact us",
        r"@",  # Email addresses

        # Legal
        r"trademark",
        r"disclaimer",
    ]

    # Page markers and headers/footers
    PAGE_MARKERS = [
        r"page \d+",
        r"^\d+$",
        r"^\d+\s*\|",
        r"\|\s*\d+$",
        r"^page \d+ of \d+",
        r"^\d+\s*/\s*\d+$",
        r"^chapter \d+$",
        r"^section \d+\.?\d*$",
        r"^\d+\.\d+$",  # Section numbers like "1.1"
    ]

    def __init__(self):
        """Initialize the document processor."""
        self.min_chunk_size = 200  # Smaller for more granular chunks
        self.max_chunk_size = 800  # Much smaller for focused content
        self.target_chunk_size = 500  # Ideal chunk size
        self._kb_manager = None  # Lazy initialization

    @property
    def kb_manager(self):
        """Lazy-load the knowledge base manager."""
        if self._kb_manager is None:
            from orchestrator.knowledge_base_manager import KnowledgeBaseManager
            self._kb_manager = KnowledgeBaseManager()
        return self._kb_manager

    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file formats.

        Args:
            file_path: Path to document file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file format not supported
        """
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self._extract_from_word(file_path)
        elif extension == '.txt':
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

    def _extract_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF file."""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError(
                "PyPDF2 is required for PDF processing. "
                "Install it with: pip install PyPDF2"
            )

        text_content = []

        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
                except Exception as e:
                    print(f"⚠️ Could not extract text from page {page_num + 1}: {e}")

        return "\n\n".join(text_content)

    def _extract_from_word(self, word_path: str) -> str:
        """Extract text from Word documents (DOCX/DOC)."""
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx is required for Word document processing. "
                "Install it with: pip install python-docx"
            )

        try:
            doc = docx.Document(word_path)
            text_content = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))

            return "\n\n".join(text_content)

        except Exception as e:
            raise ValueError(f"Error reading Word document: {e}")

    def _extract_from_txt(self, txt_path: str) -> str:
        """Extract text from TXT file."""
        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def clean_text(self, text: str) -> str:
        """Clean extracted text by aggressively removing unwanted sections.

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text with only useful content
        """
        original_length = len(text)

        # Step 1: Remove entire unwanted sections
        text = self._remove_unwanted_sections(text)

        # Step 2: Filter paragraphs to remove unwanted content
        text = self._filter_paragraphs(text)

        # Step 3: Remove page markers and headers/footers
        text = self._remove_page_markers(text)

        # Step 4: Remove repeated content
        text = self._remove_repeated_content(text)

        # Step 5: Clean up whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        text = text.strip()

        cleaned_length = len(text)
        reduction = ((original_length - cleaned_length) / original_length * 100) if original_length > 0 else 0
        print(f"📝 Cleaned text: {original_length} → {cleaned_length} chars ({reduction:.1f}% removed)")

        return text

    def _remove_unwanted_sections(self, text: str) -> str:
        """Remove entire sections that are not useful (TOC, references, etc.)."""
        # Split into paragraphs to work paragraph-by-paragraph
        paragraphs = text.split('\n\n')
        kept_paragraphs = []
        skip_until_next_section = False

        for para in paragraphs:
            para_lower = para.strip().lower()

            # Check if this paragraph is an unwanted section header
            is_unwanted_section = False
            for pattern in self.UNWANTED_SECTION_HEADERS:
                if re.match(pattern, para_lower, re.IGNORECASE):
                    is_unwanted_section = True
                    skip_until_next_section = True
                    break

            if is_unwanted_section:
                continue

            # If we're skipping, check if we hit a new major section (to stop skipping)
            if skip_until_next_section:
                # Check if this is a major content section (all caps, or chapter heading)
                if (self._is_section_header(para) and
                    not any(re.match(p, para_lower, re.IGNORECASE) for p in self.UNWANTED_SECTION_HEADERS)):
                    # This is a real content section, stop skipping
                    skip_until_next_section = False
                    kept_paragraphs.append(para)
                # Otherwise continue skipping
                continue

            # Keep this paragraph
            kept_paragraphs.append(para)

        return '\n\n'.join(kept_paragraphs)

    def _filter_paragraphs(self, text: str) -> str:
        """Filter out paragraphs containing unwanted phrases."""
        paragraphs = text.split('\n\n')
        filtered_paragraphs = []

        for para in paragraphs:
            para_lower = para.lower()
            para_stripped = para.strip()

            # Skip if contains unwanted phrases
            contains_unwanted = False
            for phrase_pattern in self.UNWANTED_PHRASES:
                if re.search(phrase_pattern, para_lower, re.IGNORECASE):
                    contains_unwanted = True
                    break

            if contains_unwanted:
                continue

            # Keep section headers even if short
            if self._is_section_header(para_stripped):
                filtered_paragraphs.append(para)
                continue

            # Skip very short paragraphs (likely formatting artifacts)
            if len(para_stripped) < 15:
                continue

            # Skip if mostly punctuation or numbers
            # Be less strict - only remove if less than 30% alphabetic
            alpha_chars = sum(c.isalpha() for c in para)
            if len(para) > 30 and alpha_chars < len(para) * 0.3:
                continue

            filtered_paragraphs.append(para)

        return '\n\n'.join(filtered_paragraphs)

    def _remove_page_markers(self, text: str) -> str:
        """Remove page numbers and headers/footers while preserving paragraph breaks."""
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            line_lower = line.lower().strip()

            # Preserve empty lines (they're paragraph separators!)
            if not line.strip():
                cleaned_lines.append(line)
                continue

            # Skip page markers
            if any(re.match(pattern, line_lower, re.IGNORECASE) for pattern in self.PAGE_MARKERS):
                continue

            # Skip very short lines (likely headers/footers) but not empty lines
            if len(line.strip()) < 10:
                continue

            # Skip lines that are just repeated characters (decorative lines)
            if line.strip() and len(set(line.strip())) <= 3:  # Like "========" or "--------"
                continue

            cleaned_lines.append(line)

        return '\n'.join(cleaned_lines)

    def _remove_repeated_content(self, text: str) -> str:
        """Remove repeated sentences or paragraphs."""
        paragraphs = text.split('\n\n')
        seen = set()
        unique_paragraphs = []

        for para in paragraphs:
            para_clean = para.strip().lower()

            if para_clean in seen:
                continue

            if len(para_clean) < 30:
                unique_paragraphs.append(para)
                continue

            seen.add(para_clean)
            unique_paragraphs.append(para)

        return '\n\n'.join(unique_paragraphs)

    def chunk_text(self, text: str, category: str = "general") -> List[Dict[str, Any]]:
        """Chunk cleaned text into small, focused knowledge base documents.

        This creates many small chunks (200-800 chars) rather than a few large ones.
        Each chunk focuses on a single topic or concept.

        Strategy:
        - Split aggressively at section headers
        - Split at paragraph boundaries when target size reached
        - Keep related bullet points together
        - Prefer smaller, focused chunks over larger ones
        """
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

        chunks = []
        current_chunk = ""
        current_title = "Safety Information"
        chunk_counter = 0

        for i, para in enumerate(paragraphs):
            # Check if this is a section header
            if self._is_section_header(para):
                # ALWAYS split at section headers
                if current_chunk.strip() and len(current_chunk) >= self.min_chunk_size * 0.5:
                    chunks.append({
                        "title": current_title,
                        "content": current_chunk.strip(),
                        "category": category
                    })
                    chunk_counter += 1
                    current_chunk = ""

                # Start new section with this header as title
                current_title = para[:100].strip()
                continue

            # Detect if we're starting a bulleted list
            is_list_item = self._is_list_item(para)
            next_is_list = False
            if i + 1 < len(paragraphs):
                next_is_list = self._is_list_item(paragraphs[i + 1])

            # Add paragraph to current chunk
            if current_chunk:
                current_chunk += "\n\n"
            current_chunk += para

            # Aggressive splitting logic
            should_split = False

            # ALWAYS split if we exceed max size
            if len(current_chunk) >= self.max_chunk_size:
                should_split = True

            # Split at target size if we're at a good break point
            elif len(current_chunk) >= self.target_chunk_size:
                # Split if:
                # 1. Next paragraph is a section header
                # 2. Next paragraph is not a list item (don't split lists)
                # 3. Current paragraph is not a list item, or list just ended
                if i + 1 < len(paragraphs):
                    next_para = paragraphs[i + 1]
                    if self._is_section_header(next_para):
                        should_split = True
                    elif not next_is_list and not is_list_item:
                        should_split = True
                    elif is_list_item and not next_is_list:
                        # End of list - good split point
                        should_split = True

            # Also split at min size if we detect topic change
            elif len(current_chunk) >= self.min_chunk_size:
                if i + 1 < len(paragraphs):
                    next_para = paragraphs[i + 1]
                    if self._is_section_header(next_para):
                        should_split = True
                    elif self._is_topic_change(para, next_para) and not is_list_item:
                        should_split = True

            if should_split:
                chunks.append({
                    "title": current_title,
                    "content": current_chunk.strip(),
                    "category": category
                })
                chunk_counter += 1
                current_chunk = ""

        # Add final chunk (be lenient with size)
        if current_chunk.strip() and len(current_chunk) >= self.min_chunk_size * 0.4:
            chunks.append({
                "title": current_title,
                "content": current_chunk.strip(),
                "category": category
            })
            chunk_counter += 1

        # If we only created 1-2 chunks from a large document, something went wrong
        # Try to split further
        if len(chunks) <= 2 and sum(len(c['content']) for c in chunks) > 2000:
            print(f"⚠️ Warning: Only created {len(chunks)} chunks. Attempting more aggressive splitting...")
            chunks = self._force_split_large_chunks(chunks, category)

        print(f"📦 Created {len(chunks)} chunks (avg size: {sum(len(c['content']) for c in chunks) // len(chunks) if chunks else 0} chars)")

        return chunks

    def _force_split_large_chunks(self, chunks: List[Dict[str, Any]], category: str) -> List[Dict[str, Any]]:
        """Force split large chunks that weren't split properly."""
        new_chunks = []

        for chunk in chunks:
            content = chunk['content']
            title = chunk['title']

            # If chunk is reasonably sized, keep it
            if len(content) <= self.max_chunk_size:
                new_chunks.append(chunk)
                continue

            # Split large chunk by paragraphs
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            sub_chunk = ""
            part_num = 1

            for para in paragraphs:
                if sub_chunk:
                    sub_chunk += "\n\n"
                sub_chunk += para

                # Split when we reach a good size
                if len(sub_chunk) >= self.target_chunk_size:
                    new_chunks.append({
                        "title": f"{title} (Part {part_num})" if part_num > 1 else title,
                        "content": sub_chunk.strip(),
                        "category": category
                    })
                    part_num += 1
                    sub_chunk = ""

            # Add remaining content
            if sub_chunk.strip():
                new_chunks.append({
                    "title": f"{title} (Part {part_num})" if part_num > 1 else title,
                    "content": sub_chunk.strip(),
                    "category": category
                })

        print(f"   Force-split into {len(new_chunks)} chunks")
        return new_chunks

    def _is_list_item(self, text: str) -> bool:
        """Check if text looks like a list item.

        Args:
            text: Text to check

        Returns:
            True if looks like a list item
        """
        # Check for common list markers
        list_markers = [
            r'^[\-\*\•]\s+',  # Dash, asterisk, bullet
            r'^\d+[\.\)]\s+',  # Numbered lists
            r'^[a-zA-Z][\.\)]\s+',  # Lettered lists
        ]

        for pattern in list_markers:
            if re.match(pattern, text.strip()):
                return True

        return False

    def _is_topic_change(self, current_para: str, next_para: str) -> bool:
        """Detect if there's a significant topic change between paragraphs.

        Args:
            current_para: Current paragraph
            next_para: Next paragraph

        Returns:
            True if appears to be a topic change
        """
        # If next paragraph starts with a strong transition or new topic indicator
        topic_change_indicators = [
            r'^(however|meanwhile|in contrast|on the other hand|alternatively)',
            r'^(step \d+|phase \d+|part \d+)',
            r'^(important|note|warning|caution|remember)',
        ]

        next_lower = next_para.lower().strip()
        for pattern in topic_change_indicators:
            if re.match(pattern, next_lower, re.IGNORECASE):
                return True

        return False

    def _is_section_header(self, text: str) -> bool:
        """Check if text looks like a section header (improved detection)."""
        text_stripped = text.strip()

        # Short and all uppercase (like "FALL PROTECTION")
        if len(text_stripped) < 100 and text_stripped.isupper() and len(text_stripped) > 5:
            return True

        # Numbered sections (like "1. Introduction" or "Chapter 1:")
        if re.match(r'^(chapter\s+)?\d+[\.:]\s+[A-Z]', text_stripped, re.IGNORECASE):
            return True

        # Section numbers (like "1.1 Safety Procedures")
        if re.match(r'^\d+\.\d+\s+[A-Z]', text_stripped):
            return True

        # Multiple capitalized words (like "Fall Protection Equipment")
        if len(text_stripped) < 100:
            words = text_stripped.split()
            if len(words) >= 2 and len(words) <= 8:
                capitalized_words = sum(1 for w in words if w and w[0].isupper())
                if capitalized_words == len(words) and not text_stripped.endswith('.'):
                    return True

        # Contains multiple all-caps words (like "OSHA FALL PROTECTION")
        if len(text_stripped) < 100 and len(re.findall(r'\b[A-Z]{2,}\b', text_stripped)) >= 2:
            return True

        return False

    def process_document(self, file_path: str, category: str = "general") -> List[Dict[str, Any]]:
        """Complete document processing pipeline.

        Args:
            file_path: Path to document file
            category: Category for the documents

        Returns:
            List of processed knowledge base documents
        """
        print(f"📄 Extracting text from: {file_path}")
        raw_text = self.extract_text_from_file(file_path)

        print(f"✨ Cleaning text ({len(raw_text)} characters)")
        cleaned_text = self.clean_text(raw_text)

        print(f"📦 Chunking into documents ({len(cleaned_text)} characters)")
        chunks = self.chunk_text(cleaned_text, category)

        print(f"✅ Created {len(chunks)} knowledge base documents")

        return chunks

    def save_chunks_to_file(self, chunks: List[Dict[str, Any]], output_path: str):
        """Save processed chunks to a Python file."""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# Auto-generated knowledge base from document\n\n")
            f.write("KNOWLEDGE_BASE = [\n")

            for chunk in chunks:
                f.write("    {\n")
                f.write(f'        "title": """{chunk["title"]}""",\n')
                f.write(f'        "content": """{chunk["content"]}""",\n')
                f.write(f'        "category": "{chunk["category"]}"\n')
                f.write("    },\n")

            f.write("]\n")

        print(f"💾 Saved to: {output_path}")

    def append_chunks_to_base_file(self, chunks: List[Dict[str, Any]], base_file_path: str):
        """Append processed chunks to an existing knowledge base file.

        This method now works with JSON files (.json) instead of Python files (.py).
        For backward compatibility, it also detects .py files and converts them.

        Args:
            chunks: List of document chunks to append
            base_file_path: Path to the base knowledge file (e.g., fall_base.json or fall_base.py)
        """
        if not chunks:
            return

        from pathlib import Path
        base_path = Path(base_file_path)

        # Determine category from filename
        # e.g., "fall_base.json" -> "fall" or "fall_base.py" -> "fall"
        category = base_path.stem.replace("_base", "")

        # If it's a .py file, convert to .json path
        if base_path.suffix == ".py":
            print(f"⚠️ Warning: Python knowledge base files are deprecated. Converting to JSON...")
            base_path = base_path.with_suffix(".json")

        # Use the knowledge base manager to append
        try:
            self.kb_manager.append_to_knowledge_base(category, chunks)
        except Exception as e:
            print(f"❌ Error appending chunks: {e}")
            raise
